from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "reports/LegacyVault_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr[index], header, bold=True)
        set_cell_shading(hdr[index], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level < 3 else 120)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return paragraph


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("LegacyVault Project Report | Page ")
    run.font.size = Pt(9)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def build_report():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("LEGACYVAULT\nBlockchain Inheritance and Smart Wallet System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subrun = subtitle.add_run("A Project Report Submitted in Partial Fulfillment of Academic Requirements")
    subrun.font.name = "Calibri"
    subrun.font.size = Pt(13)

    add_table(
        document,
        ["Field", "Details"],
        [
            ["Student Name", "[Student Name]"],
            ["Roll Number", "[Roll Number]"],
            ["Course / Branch", "[Course / Branch]"],
            ["Semester", "[Semester]"],
            ["Guide / Faculty", "[Guide Name]"],
            ["College / University", "[College / University Name]"],
            ["Academic Year", "[Academic Year]"],
        ],
        widths=[2.0, 4.2],
    )

    document.add_page_break()

    add_heading(document, "Certificate", 1)
    add_body(
        document,
        "This is to certify that the project report entitled \"LegacyVault: Blockchain Inheritance and Smart Wallet System\" "
        "has been prepared by [Student Name] under the guidance of [Guide Name]. The work presented in this report is "
        "submitted as part of the academic requirements of [Course / Branch] at [College / University Name]."
    )
    add_body(document, "Signature of Guide: __________________________")
    add_body(document, "Signature of Student: ________________________")
    add_body(document, "Date: ________________________")

    add_heading(document, "Declaration", 1)
    add_body(
        document,
        "I hereby declare that this project report is based on my own work carried out for the LegacyVault project. "
        "The project uses blockchain smart contracts, a backend oracle server, and a React-based frontend to demonstrate "
        "secure inheritance, recovery, and daily vault payment features. Any external libraries and frameworks used in "
        "the project are acknowledged in the relevant sections."
    )

    add_heading(document, "Acknowledgement", 1)
    add_body(
        document,
        "I would like to express my sincere gratitude to my project guide, faculty members, and institution for providing "
        "the support and environment required to complete this project. I also acknowledge the open-source tools and "
        "developer communities that made it possible to build and test the blockchain, backend, and frontend components."
    )

    document.add_page_break()

    add_heading(document, "Abstract", 1)
    add_body(
        document,
        "LegacyVault is a blockchain-based inheritance and smart wallet system designed to solve a critical problem in "
        "self-custody finance: what happens to crypto assets if the owner becomes inactive, loses access to a wallet, "
        "or wants to safely use funds in daily life. The project combines an upgradeable Ethereum smart contract, a "
        "Node.js oracle backend, and a React frontend. It supports vault deposits, heir configuration, inactivity-based "
        "inheritance, heartbeat updates, biometric/passkey recovery, guardian-assisted recovery, daily spending limits, "
        "and QR-based payments."
    )
    add_body(
        document,
        "The system is implemented as a Sepolia testnet demonstration. It emphasizes wallet ownership, smart contract "
        "verification, read-only feature detection, user-friendly recovery workflows, and a modern web application "
        "experience. The project does not attempt to recover a user's private key. Instead, it offers secure transfer "
        "and recovery mechanisms built around contract-controlled assets and signed authorization flows."
    )

    add_heading(document, "Table of Contents", 1)
    toc = [
        "1. Introduction",
        "2. Problem Statement",
        "3. Objectives",
        "4. Scope of the Project",
        "5. Existing System and Proposed System",
        "6. Technology Stack",
        "7. System Architecture",
        "8. Folder Structure",
        "9. Smart Contract Design",
        "10. Backend Oracle Server",
        "11. Frontend Web Application",
        "12. Functional Modules",
        "13. Recovery System",
        "14. Payments and QR Wallet Feature",
        "15. Database Design",
        "16. Security Analysis",
        "17. Testing and Validation",
        "18. Limitations",
        "19. Future Scope",
        "20. Conclusion",
        "21. References",
    ]
    add_numbered(document, toc)

    document.add_page_break()

    add_heading(document, "1. Introduction", 1)
    add_body(
        document,
        "The growth of decentralized finance has given users direct control over their assets. However, self-custody also "
        "creates practical risks. If a user loses a private key, forgets wallet access, becomes inactive, or passes away, "
        "crypto assets can become permanently inaccessible. Traditional financial systems provide nominee and inheritance "
        "mechanisms, but blockchain wallets usually do not include such protections by default."
    )
    add_body(
        document,
        "LegacyVault addresses this gap by providing a smart contract vault where users can deposit ETH, designate an heir, "
        "send periodic heartbeat transactions, recover assets through passkey or guardian verification, and use vault funds "
        "for daily payments under a configurable daily limit."
    )

    add_heading(document, "2. Problem Statement", 1)
    add_body(
        document,
        "In a self-custody wallet system, there is no central authority that can reset access or release funds to heirs. "
        "This creates three major problems: inheritance risk, wallet loss risk, and usability risk. A user may want to "
        "store wealth securely but also use it in daily life without constantly transferring funds out of a vault."
    )
    add_bullets(
        document,
        [
            "Crypto assets can be lost permanently if the owner loses private keys.",
            "Heirs may not be able to access assets after the owner's long-term inactivity.",
            "Basic vault systems often lock funds but do not support daily spending.",
            "Recovery systems must avoid weak centralized password-reset models.",
            "A frontend must be simple enough for non-expert users to operate safely.",
        ],
    )

    add_heading(document, "3. Objectives", 1)
    add_bullets(
        document,
        [
            "Design a blockchain vault for ETH deposits and inheritance configuration.",
            "Allow the owner to set an heir and inactivity period.",
            "Provide a heartbeat mechanism to prove that the owner is still active.",
            "Support biometric/passkey recovery through a backend oracle signing flow.",
            "Support guardian-based recovery as a demo social recovery mechanism.",
            "Enable daily wallet-style payments from vault funds using spending limits.",
            "Provide QR-based payment request and scan functionality in the frontend.",
            "Create a premium web dashboard for managing vault, recovery, and payments.",
        ],
    )

    add_heading(document, "4. Scope of the Project", 1)
    add_body(
        document,
        "The current project is a testnet-level decentralized application built for demonstration and academic evaluation. "
        "It focuses on core workflows and compatibility with the deployed Sepolia contract. It does not deploy a production "
        "mainnet financial product."
    )
    add_table(
        document,
        ["Included", "Not Included"],
        [
            ["ETH vault deposit and management", "Mainnet deployment"],
            ["Heir setup and inactivity rules", "Legal inheritance enforcement"],
            ["Passkey and guardian recovery demo", "Private key recovery"],
            ["Daily ETH spending from vault", "Fiat payment rails such as UPI"],
            ["QR payment links and scanner fallback", "Native mobile app"],
        ],
        widths=[3.1, 3.1],
    )

    add_heading(document, "5. Existing System and Proposed System", 1)
    add_heading(document, "5.1 Existing System", 2)
    add_body(
        document,
        "Most standard wallets depend entirely on the user's seed phrase or private key. If the key is lost, the wallet "
        "provider cannot recover funds. Some custodial exchanges provide account recovery, but they require users to "
        "give control of funds to a third party."
    )
    add_heading(document, "5.2 Proposed System", 2)
    add_body(
        document,
        "LegacyVault proposes a self-custody smart contract wallet where assets remain controlled by on-chain rules. "
        "The user can define heirs, maintain activity with heartbeat, recover through passkey or guardian verification, "
        "and perform daily spending through a contract-level limit."
    )

    add_heading(document, "6. Technology Stack", 1)
    add_table(
        document,
        ["Layer", "Technology"],
        [
            ["Blockchain", "Solidity, Foundry, OpenZeppelin Upgradeable Contracts, Sepolia Testnet"],
            ["Smart Contract Pattern", "UUPS upgradeable proxy pattern"],
            ["Backend", "Node.js, Express.js, ethers.js, CORS, dotenv, JSON Web Token"],
            ["Database", "MongoDB with Mongoose models"],
            ["Frontend", "React, Vite, Wagmi, Viem, RainbowKit, CSS/Tailwind utilities"],
            ["Authentication / Recovery", "WebAuthn passkeys, JWT oracle flow, guardian signatures"],
            ["Payments", "V3 spend function, QR code generation, native browser QR scan fallback"],
        ],
        widths=[2.0, 4.2],
    )

    add_heading(document, "7. System Architecture", 1)
    add_body(
        document,
        "The project follows a three-layer architecture. The blockchain layer stores vault rules and funds. The backend "
        "oracle verifies biometric or guardian approval flows and signs recovery authorization messages. The frontend "
        "provides the user interface for wallet connection, contract calls, recovery actions, and QR payments."
    )
    add_code_block(
        document,
        "User Browser / React UI\n"
        "  |-- Wallet connection through RainbowKit/Wagmi\n"
        "  |-- Contract reads and writes through Viem\n"
        "  |-- Oracle API calls for passkey and guardian recovery\n"
        "        |\n"
        "        v\n"
        "Node.js Oracle Backend\n"
        "  |-- WebAuthn registration and verification\n"
        "  |-- Guardian setup and approval tracking\n"
        "  |-- JWT session generation\n"
        "  |-- Oracle signature generation\n"
        "        |\n"
        "        v\n"
        "Ethereum Sepolia Smart Contract\n"
        "  |-- Vault deposit, heir, inactivity, heartbeat\n"
        "  |-- Recovery signature verification\n"
        "  |-- Daily spending limit and spend function",
    )

    add_heading(document, "8. Folder Structure", 1)
    add_body(document, "The repository is organized into smart contract, backend, and frontend areas.")
    add_code_block(
        document,
        "legacy-vault/\n"
        "  src/                         Solidity smart contracts\n"
        "  script/                      Foundry deployment and upgrade scripts\n"
        "  test/                        Solidity tests\n"
        "  Server/                      Backend models and biometric routes\n"
        "  oracle-server.js             Main oracle/backend server\n"
        "  package.json                 Root backend dependencies and scripts\n"
        "  legacy-vault-ui/             React frontend application\n"
        "    src/components/            Dashboard, vault, recovery, payment components\n"
        "    src/abis/                  Frontend contract ABI\n"
        "    src/config/                Contract address and environment config\n"
        "    package.json               Frontend dependencies and scripts",
    )

    add_heading(document, "9. Smart Contract Design", 1)
    add_body(
        document,
        "The smart contract layer is based on upgradeable Solidity contracts. The core contract stores vault data per user, "
        "including balance, inactivity period, last heartbeat, and heir address. Later versions extend the contract with "
        "nonce-based biometric recovery and daily spending functions."
    )
    add_table(
        document,
        ["Contract Version", "Main Functions"],
        [
            ["V1", "deposit, setHeir, setInactivityPeriod, heartbeat, triggerInheritance, recoverWithBiometric(address,bytes)"],
            ["V2", "usedNonces mapping and recoverWithBiometric(address,bytes,uint256,uint256)"],
            ["V3", "dailyLimit, spentToday, setDailyLimit, spend, updated version detection"],
        ],
        widths=[1.4, 4.8],
    )

    add_heading(document, "9.1 Vault Data Structure", 2)
    add_code_block(
        document,
        "struct Vault {\n"
        "    uint256 balance;\n"
        "    uint256 inactivityPeriod;\n"
        "    uint256 lastHeartbeat;\n"
        "    address heir;\n"
        "}",
    )
    add_heading(document, "9.2 Important Contract Functions", 2)
    add_table(
        document,
        ["Function", "Purpose"],
        [
            ["deposit()", "Allows a user to deposit ETH into their own vault."],
            ["setHeir(address)", "Stores the heir wallet address for inheritance."],
            ["setInactivityPeriod(uint256)", "Sets the time period after which inheritance can be triggered."],
            ["heartbeat()", "Updates the last active timestamp for the owner."],
            ["triggerInheritance(address)", "Transfers vault funds to the configured heir after inactivity."],
            ["recoverWithBiometric(...)", "Allows oracle-approved recovery to a replacement wallet."],
            ["setDailyLimit(uint256)", "Sets maximum vault spending allowed per day."],
            ["spend(address,uint256)", "Sends ETH from vault to a recipient under the daily limit."],
        ],
        widths=[2.3, 3.9],
    )

    add_heading(document, "10. Backend Oracle Server", 1)
    add_body(
        document,
        "The backend oracle server is responsible for off-chain verification and recovery authorization. It does not hold "
        "user funds. Its main job is to verify passkey or guardian recovery conditions and produce a signature that the "
        "smart contract can verify."
    )
    add_table(
        document,
        ["Endpoint", "Description"],
        [
            ["/register/start", "Starts WebAuthn passkey registration."],
            ["/register/finish", "Completes passkey registration."],
            ["/verify/start", "Starts passkey verification for recovery."],
            ["/verify/finish", "Verifies passkey and returns a JWT token."],
            ["/guardians/save", "Stores guardian wallets after owner signature."],
            ["/guardian-recovery/start", "Creates or loads a guardian recovery request."],
            ["/guardian-recovery/approve", "Records a guardian's signed approval."],
            ["/guardian-recovery/token", "Returns JWT when guardian threshold is met."],
            ["/sign-recovery", "Signs recovery payload for contract verification."],
        ],
        widths=[2.5, 3.7],
    )

    add_heading(document, "10.1 Oracle Signature Logic", 2)
    add_body(
        document,
        "For V2/V3 recovery, the backend signs a hash containing the lost wallet, caller/replacement wallet, chain ID, "
        "contract address, nonce, and expiry. This must match the Solidity verification logic."
    )
    add_code_block(
        document,
        "hash = keccak256(abi.encodePacked(\n"
        "    user,\n"
        "    caller,\n"
        "    chainId,\n"
        "    contractAddress,\n"
        "    nonce,\n"
        "    expiry\n"
        "))",
    )

    add_heading(document, "11. Frontend Web Application", 1)
    add_body(
        document,
        "The frontend is a Vite React application. It uses Wagmi and Viem for contract interaction and RainbowKit for "
        "wallet connection. The frontend is organized as a web app with a sidebar and dedicated pages for dashboard, "
        "vault operations, heir setup, recovery, payments, and activity timeline."
    )
    add_table(
        document,
        ["Route", "Purpose"],
        [
            ["/app", "Dashboard summary and vault health overview."],
            ["/app/vault", "Deposit, biometric registration, inactivity, and heartbeat actions."],
            ["/app/heir", "Set heir and trigger inheritance flow."],
            ["/app/recovery", "Passkey and guardian recovery center."],
            ["/app/recovery/guardian-sign", "Guardian signing page for shared recovery links."],
            ["/app/pay", "Daily wallet payments, QR scan, receive QR, payment links."],
            ["/app/activity", "Local activity timeline."],
        ],
        widths=[2.4, 3.8],
    )

    add_heading(document, "12. Functional Modules", 1)
    add_heading(document, "12.1 Wallet Connection", 2)
    add_body(
        document,
        "Users connect wallets through RainbowKit. The app checks the connected chain and warns users if they are not on "
        "Sepolia. Wallet connection is required for contract writes."
    )
    add_heading(document, "12.2 Deposit and Vault Management", 2)
    add_body(
        document,
        "The vault page allows ETH deposit, biometric registration, heartbeat, and inactivity configuration. Deposit is "
        "a smart contract function and does not require biometric registration by contract design."
    )
    add_heading(document, "12.3 Heir Management", 2)
    add_body(
        document,
        "The heir page allows the owner to set the heir wallet. The heir or any caller can trigger inheritance after the "
        "owner has been inactive beyond the configured period, if the contract conditions are satisfied."
    )
    add_heading(document, "12.4 Activity Timeline", 2)
    add_body(
        document,
        "The frontend stores recent actions in localStorage. This provides a simple activity timeline without requiring "
        "additional backend schema changes."
    )

    add_heading(document, "13. Recovery System", 1)
    add_heading(document, "13.1 Passkey Recovery", 2)
    add_body(
        document,
        "Passkey recovery uses WebAuthn. The user must have registered a passkey earlier. During recovery, the browser "
        "prompts the user for biometric/passkey verification. If verification succeeds, the backend returns a JWT token "
        "that permits an oracle recovery signature request."
    )
    add_numbered(
        document,
        [
            "Enter the lost wallet address.",
            "Start passkey verification.",
            "Complete browser biometric/passkey prompt.",
            "Receive JWT from backend after successful verification.",
            "Connect replacement wallet.",
            "Request oracle signature and submit recoverWithBiometric transaction.",
        ],
    )

    add_heading(document, "13.2 Guardian Recovery", 2)
    add_body(
        document,
        "Guardian recovery is implemented as a demo social recovery flow. The owner saves guardian addresses and a threshold. "
        "During recovery, a request is created and guardians sign an off-chain approval message. Once enough approvals are "
        "collected, the backend issues a JWT token for recovery signing."
    )
    add_numbered(
        document,
        [
            "Owner saves guardian wallet addresses and threshold.",
            "Recovering user enters lost wallet and replacement wallet.",
            "Application creates a guardian recovery request.",
            "Guardian signs either on the same device for demo or through a shared request link.",
            "Backend records approvals in memory for the demo session.",
            "When threshold is met, the backend returns a recovery JWT.",
            "Replacement wallet submits the final recovery transaction.",
        ],
    )
    add_body(
        document,
        "Important note: A website cannot automatically open a wallet popup on another person's device only from an address. "
        "The practical Web3-safe approach is to generate a request link that the guardian opens on their own device."
    )

    add_heading(document, "14. Payments and QR Wallet Feature", 1)
    add_body(
        document,
        "LegacyVault includes a daily wallet feature so funds stored in the vault are not idle. The user can set a daily "
        "limit and pay recipients from vault funds. This makes the vault more useful for regular transactions while still "
        "maintaining inheritance and recovery safety."
    )
    add_table(
        document,
        ["Payment Feature", "Description"],
        [
            ["Pay by Address", "Send ETH to a receiver wallet/public key using spend(address,uint256)."],
            ["Daily Limit", "Restricts how much can be spent in one day."],
            ["Recent Payments", "Stores recent recipients in localStorage for faster reuse."],
            ["Payment Link", "Creates /app/pay?to=...&amount=... links for payment requests."],
            ["Receive QR", "Generates a QR code for the connected wallet and optional amount."],
            ["QR Scanner", "Uses native browser BarcodeDetector when available; paste fallback exists."],
        ],
        widths=[2.0, 4.2],
    )
    add_body(
        document,
        "Current limitation: the receive QR points to the receiver's wallet address. Direct deposit into another user's "
        "contract vault would require a future smart contract function such as depositFor(address)."
    )

    add_heading(document, "15. Database Design", 1)
    add_body(
        document,
        "The backend uses MongoDB through Mongoose. The User model stores wallet address, registered authenticators for "
        "passkeys, guardians, and guardian threshold."
    )
    add_code_block(
        document,
        "User\n"
        "  walletAddress: string\n"
        "  authenticators: array\n"
        "    credentialID\n"
        "    credentialPublicKey\n"
        "    counter\n"
        "  guardians: array\n"
        "    address\n"
        "  guardianThreshold: number",
    )

    add_heading(document, "16. Security Analysis", 1)
    add_bullets(
        document,
        [
            "The smart contract keeps custody of vault funds rather than the backend.",
            "Recovery signatures include chain ID and contract address to reduce replay risk.",
            "V2/V3 recovery uses nonce and expiry to prevent reuse of old recovery signatures.",
            "Guardian setup requires the owner to sign the guardian configuration message.",
            "Guardian approval requires signatures from configured guardian wallets.",
            "The frontend checks V2/V3 feature availability before showing recovery and payment actions.",
            "Daily spending limits reduce risk from accidental or excessive spending.",
        ],
    )
    add_heading(document, "16.1 Security Limitations", 2)
    add_bullets(
        document,
        [
            "Guardian approvals are stored in memory for the demo and reset after backend restart.",
            "The oracle private key must be protected carefully in a real deployment.",
            "The project is designed for Sepolia demonstration, not production mainnet use.",
            "Legal inheritance validation is outside the scope of the smart contract demo.",
            "Browser QR scanner support varies by browser; fallback input is required.",
        ],
    )

    add_heading(document, "17. Testing and Validation", 1)
    add_body(
        document,
        "The application was validated through linting, build checks, contract compatibility review, frontend runtime testing, "
        "and manual flow testing. The frontend build confirms that routes, components, and dependencies compile successfully."
    )
    add_table(
        document,
        ["Test Case", "Expected Result"],
        [
            ["Backend startup", "Oracle server starts on configured port and CORS allows frontend origin."],
            ["Frontend startup", "Vite dev server opens and app shell loads."],
            ["Wallet connect", "Wallet modal opens and connected address appears in header."],
            ["Sepolia check", "Warning appears if wrong network is selected."],
            ["Deposit", "Vault deposit transaction is submitted and balance refreshes."],
            ["Set heir", "Heir address is stored on-chain."],
            ["Heartbeat", "Last heartbeat timestamp updates."],
            ["Passkey recovery", "Registered passkey verifies and returns JWT."],
            ["Guardian recovery", "Guardian signatures reach threshold and recovery token is generated."],
            ["Payments", "V3 spend function sends ETH under daily limit."],
            ["QR payment", "Payment link/QR loads receiver and amount into payment form."],
        ],
        widths=[2.1, 4.1],
    )

    add_heading(document, "18. Limitations", 1)
    add_bullets(
        document,
        [
            "The project currently targets Sepolia testnet only.",
            "The backend oracle is centralized in the demo architecture.",
            "Guardian recovery approvals are not persisted permanently in the current demo.",
            "Receive QR sends to wallet address, not directly into another user's vault balance.",
            "The app does not support ERC-20 tokens yet.",
            "Native browser QR scanning may not work in every browser.",
            "Production deployment would require audits, monitoring, and stronger key management.",
        ],
    )

    add_heading(document, "19. Future Scope", 1)
    add_bullets(
        document,
        [
            "Add depositFor(address) to allow direct QR deposits into another user's vault.",
            "Persist guardian recovery requests and approvals in the database.",
            "Add email, push, or messaging notifications for guardian signing links.",
            "Support ERC-20 tokens and stablecoins for daily wallet payments.",
            "Add ENS or contact aliases for easier wallet address management.",
            "Add mobile-first QR scanning with a dedicated QR scanning package.",
            "Add multi-oracle or decentralized oracle approval for stronger recovery trust.",
            "Add formal smart contract security audit before mainnet deployment.",
            "Add production monitoring, logs, and rate limits to backend APIs.",
        ],
    )

    add_heading(document, "20. Conclusion", 1)
    add_body(
        document,
        "LegacyVault demonstrates how blockchain smart contracts can be used to combine inheritance planning, account "
        "recovery, and daily wallet utility. The project addresses a real-world problem in self-custody finance by "
        "allowing users to store funds securely while still enabling heir transfer, heartbeat-based activity checks, "
        "passkey recovery, guardian recovery, and daily QR payments."
    )
    add_body(
        document,
        "The project successfully integrates Solidity smart contracts, a Node.js oracle backend, MongoDB data storage, "
        "and a React web application. Although it remains a testnet demonstration, it provides a strong foundation for "
        "future development into a more production-ready smart wallet and inheritance platform."
    )

    add_heading(document, "21. References", 1)
    add_bullets(
        document,
        [
            "Solidity Documentation: https://docs.soliditylang.org/",
            "Foundry Book: https://book.getfoundry.sh/",
            "OpenZeppelin Contracts: https://docs.openzeppelin.com/contracts/",
            "Ethereum Sepolia Testnet Documentation",
            "React Documentation: https://react.dev/",
            "Vite Documentation: https://vite.dev/",
            "Wagmi Documentation: https://wagmi.sh/",
            "Viem Documentation: https://viem.sh/",
            "RainbowKit Documentation: https://www.rainbowkit.com/",
            "WebAuthn / Passkeys Overview: https://webauthn.guide/",
        ],
    )

    add_heading(document, "Appendix A: Environment Variables", 1)
    add_table(
        document,
        ["Location", "Important Variables"],
        [
            ["Backend root .env", "ORACLE_PORT, MONGO_URI, JWT_SECRET, ORACLE_PRIVATE_KEY, CONTRACT_ADDRESS, CHAIN_ID, FRONTEND_ORIGINS"],
            ["Frontend .env", "VITE_VAULT_ADDRESS, VITE_CHAIN_ID, VITE_ORACLE_BASE_URL, VITE_SEPOLIA_RPC_URL, VITE_WALLETCONNECT_PROJECT_ID"],
        ],
        widths=[2.0, 4.2],
    )

    add_heading(document, "Appendix B: User Operation Checklist", 1)
    add_numbered(
        document,
        [
            "Start the oracle backend.",
            "Start the frontend development server.",
            "Connect wallet and switch to Sepolia.",
            "Deposit ETH into vault.",
            "Register passkey if recovery by passkey is required.",
            "Set heir wallet address.",
            "Set inactivity period.",
            "Send heartbeat periodically.",
            "Configure daily spending limit.",
            "Pay by wallet address or QR code.",
            "Create receive QR for payment request.",
            "Configure guardians for social recovery.",
            "Test guardian signing link from another wallet.",
            "Verify activity timeline and vault health score.",
        ],
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    build_report()
